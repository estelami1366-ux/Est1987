# -*- coding: utf-8 -*-
"""Honest Print Center verification report. Does not claim physical print success."""
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

OUT = Path("/workspace/PRINT_CENTER_FINAL_VERIFICATION_REPORT.docx")
COPIES = [
    Path("/workspace/deliveries/Reports/PRINT_CENTER_FINAL_VERIFICATION_REPORT.docx"),
]

STATUS = "BLOCKED — REAL WINDOWS PRINTER VERIFICATION REQUIRED"


def set_rtl(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        from docx.oxml import OxmlElement
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
    bidi.set(qn("w:val"), "1")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_h(doc, text, size=16, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "Tahoma"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    if color:
        run.font.color.rgb = RGBColor(*color)
    set_rtl(p)
    return p


def add_p(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Tahoma"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    set_rtl(p)
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        run.font.size = Pt(11)
        run.font.name = "Tahoma"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
        set_rtl(p)


def main():
    doc = Document()
    for s in doc.sections:
        s.right_margin = Cm(2)
        s.left_margin = Cm(2)

    add_h(doc, "گزارش نهایی راستی‌آزمایی مرکز پرینت سیرمان", 20)
    add_p(doc, "نسخه برنامه: ۱۴۰۵.۵.۲۵β  /  1405.5.25β  /  اسمبلی 1405.5.25.2")
    add_p(doc, "تاریخ شمسی گزارش: ۱۴۰۵/۰۵/۲۵")
    add_p(doc, "محیط این بررسی: عامل ابری لینوکس — بدون WebView2، بدون اسپولر ویندوز، بدون چاپگر فیزیکی.")
    add_h(doc, STATUS, 16, color=(176, 32, 32))
    add_p(
        doc,
        "این گزارش چاپ روی کاغذ را تأیید نمی‌کند. ادعاهای FIXED / PRINT WORKS / VERIFIED برای خروجی فیزیکی ممنوع است تا وقتی یک رایانه ویندوز فروشگاه با همین Sirman.exe و چاپگر واقعی یک برگه چاپ کند.",
        bold=True,
    )

    add_h(doc, "۱) معماری چاپ بعد از اصلاح")
    add_p(doc, "مسیر واحد چاپ این است:")
    add_bullets(
        doc,
        [
            "UI / JS (مرکز پرینت، چاپ سریع فاکتور، چاپ فروش) → printEnginePrintHtml / PrintDocument",
            "Host Object: sirmanHost.PrintDocument یا sirmanHost.PrintHtml",
            "MainForm.EnqueueHtmlPrint",
            "WindowsPrintHost: WebView2 مخفی → CoreWebView2.PrintAsync → اسپولر چاپ ویندوز",
        ],
    )
    add_p(doc, "وضعیت کار چاپ:")
    add_bullets(
        doc,
        [
            "PRINT_REQUESTED — کار در صف برنامه ثبت شد (هنوز کاغذ نیست)",
            "PRINTING — سند در حال ارسال به اسپولر",
            "PRINT_SUBMITTED — اسپولر ویندوز سند را پذیرفت (موفقیت ارسال، نه تأیید برگه روی میز)",
            "PRINT_FAILED — شکست صادقانه با errorCode و errorMessage",
        ],
    )
    add_p(
        doc,
        "برنامه هرگز PRINT_COMPLETED به‌معنی «کاغذ بیرون آمد» اعلام نمی‌کند. حتی PRINT_SUBMITTED فقط یعنی ویندوز کار را به صف چاپگر داد.",
    )

    add_h(doc, "۲) علت ریشه‌ای خرابی چاپ قبلی")
    add_p(doc, "چاپ قبلی واقعی نبود. چند لایه موفقیت جعلی روی هم نشسته بود:")
    add_bullets(
        doc,
        [
            "UI → JS printEnginePrintHtml → Host.PrintHtml → نوشتن فایل موقت HTML → Process.Start با فعل شل print یا printto → بلافاصله {\"ok\":true}.",
            "ویندوز فقط «چاپ را شروع کن» را می‌پذیرفت. نتیجه اسپولر، کاغذ، جهت و اندازه خوانده نمی‌شد.",
            "کاغذ (A4/A5/80mm) و جهت سند دور ریخته می‌شد.",
            "GetPrinters اگر فهرست خالی بود «Microsoft Print to PDF» جعلی می‌ساخت.",
            "JS هم در نبود میزبان چاپگر جعلی می‌ساخت (printEngineFallbackPrinters).",
            "pcDoPrint پیش‌نمایش نمونه را چاپ می‌کرد، نه فاکتور زنده.",
            "printSaleDoc با window.open + window.print از موتور مرکزی دور می‌زد.",
        ],
    )
    add_p(
        doc,
        "نقطه شکست دقیق: SirmanHostObject.PrintHtml بعد از Process.Start بدون انتظار برای PrintAsync / وضعیت اسپولر، ok:true برمی‌گرداند. کاربر «ارسال به چاپگر» می‌دید ولی برگه چاپ نمی‌شد.",
        bold=True,
    )

    add_h(doc, "۳) فایل‌ها و متدهای درگیر")
    add_bullets(
        doc,
        [
            "desktop/Sirman.Desktop/WindowsPrintHost.cs — موتور جدید: ListPrinters، Enqueue، RunAsync، PrintAsync، ResolvePrinter، ApplyPaper",
            "desktop/Sirman.Desktop/MainForm.cs — ListPrintersJson، GetPrintJobJson، EnqueueHtmlPrint",
            "desktop/Sirman.Desktop/SirmanHostObject.cs — GetPrinters، GetPrintJob، PrintHtml، PrintDocument (بدون printto / Process.Start)",
            "desktop/Sirman.Core/Security/PermissionCatalog.cs — PrintHtml / PrintDocument / GetPrintJob → Print.Use",
            "Sirman_Final.html — printEnginePrintHtml، printEngineListPrinters، pcDoPrint، pcDoPdf، pcDoRetry، openFreshPrintWindow، printInv، printSvSingle، printSaleDoc",
            "لاگ کار چاپ: %AppData%\\Sirman\\print\\print-jobs.jsonl",
        ],
    )

    add_h(doc, "۴) کشف و انتخاب چاپگر")
    add_bullets(
        doc,
        [
            "منبع فهرست: PrinterSettings.InstalledPrinters ویندوز. چاپگر جعلی ساخته نمی‌شود.",
            "اگر فهرست خالی باشد: NO_PRINTER — «چاپگری نصب نیست.»",
            "اگر نام خالی باشد: چاپگر پیش‌فرض معتبر. اگر پیش‌فرض نباشد: NO_DEFAULT_PRINTER.",
            "اگر نام انتخاب‌شده در فهرست نباشد: PRINTER_NOT_FOUND.",
            "اگر IsValid=false یا PrintAsync=PrinterUnavailable: PRINTER_UNAVAILABLE با متن انگلیسی «Printer is unavailable.»",
            "بدون Sirman.exe: NO_HOST — چاپ واقعی فقط در exe.",
        ],
    )

    add_h(doc, "۵) جریان کار چاپ")
    add_bullets(
        doc,
        [
            "کاربر چاپ می‌زند → JS HTML زنده + InvoiceId/documentId را می‌فرستد.",
            "Host کار را Enqueue می‌کند و PRINT_REQUESTED + printJobId برمی‌گرداند.",
            "UI وضعیت «در حال چاپ» نشان می‌دهد و GetPrintJob را می‌پرسد.",
            "رشته UI ویندوز سند را در WebView2 مخفی بار می‌کند و PrintAsync می‌زند.",
            "Succeeded → PRINT_SUBMITTED. در غیر این صورت PRINT_FAILED با کد مشخص.",
            "تاریخچه مرکز پرینت همان وضعیت را ثبت می‌کند. پیش‌فرض تاریخچه PRINT_FAILED است نه ok.",
        ],
    )

    add_h(doc, "۶) خطاهای صادقانه")
    add_bullets(
        doc,
        [
            "NO_HOST — برنامه در مرورگر/HTML-only است، نه Sirman.exe",
            "NO_UI — پنجره برنامه برای چاپ آماده نیست",
            "NO_DOCUMENT — سند زنده‌ای برای چاپ نیست (نمونه پیش‌نمایش چاپ نمی‌شود)",
            "NO_PRINTER — چاپگری نصب نیست",
            "NO_DEFAULT_PRINTER — چاپگر پیش‌فرض تنظیم نشده",
            "PRINTER_NOT_FOUND — چاپگر انتخاب‌شده پیدا نشد",
            "PRINTER_UNAVAILABLE — چاپگر در دسترس نیست",
            "NO_PDF_PRINTER — چاپگر PDF ویندوز نصب نیست",
            "WEBVIEW / NAVIGATION / HOST_ERROR / BAD_RESPONSE / TIMEOUT / PRINT_FAILED",
        ],
    )

    add_h(doc, "۷) فاکتور، RTL، A4 — فقط در سطح کد")
    add_p(
        doc,
        "در کد: HTML فاکتور dir=rtl است، فونت Tahoma، شناسه چاپ InvoiceId است نه فقط شماره روی برگه، و اندازه صفحه A4/A5/80mm/برچسب به CoreWebView2PrintSettings داده می‌شود. این‌ها در لینوکس روی کاغذ دیده نشد. تأیید چیدمان فارسی روی برگه A4 فقط روی ویندوز فروشگاه ممکن است.",
    )

    add_h(doc, "۸) آفلاین / بدون چاپگر — فقط در سطح کد")
    add_p(
        doc,
        "اگر چاپگر نباشد یا آفلاین باشد، کد موفقیت جعلی برنمی‌گرداند. فهرست خالی PDF جعلی نمی‌سازد. این رفتار با تست Node و خواندن سورس C# بررسی شد. رفتار واقعی اسپولر آفلاین روی این محیط اجرا نشد.",
    )

    add_h(doc, "۹) ری‌استارت برنامه")
    add_p(
        doc,
        "ری‌استارت Sirman.exe روی این محیط تست نشد. کار چاپ در حافظه پروسه است؛ بعد از بستن exe وضعیت در حال انجام از بین می‌رود. لاگ JSONL روی دیسک می‌ماند ولی صف چاپ ویندوز مال خود ویندوز است.",
    )

    add_h(doc, "۱۰) رگرسیون")
    add_p(doc, "تست HTML: node test_laegh.js Sirman_Final.html → ۵۳۷ موفق / ۰ ناموفق.")
    add_p(doc, "تست C#: dotnet test desktop/Sirman.Core.Tests → ۹۲ موفق / ۰ ناموفق.")
    add_p(doc, "این اعداد فقط منطق برنامه و سورس چاپ را پوشش می‌دهند، نه برگه کاغذ.")
    add_p(doc, "چاپ نباید فاکتور، موجودی کالا/قطعه، یا مانده حساب را عوض کند. تست شبیه‌سازی این را چک می‌کند.")
    add_p(doc, "شماره‌گذاری فاکتور عوض نشده است. هویت حذف همچنان InvoiceId است.")

    add_h(doc, "۱۱) ریسک باقی‌مانده")
    add_bullets(
        doc,
        [
            "EXE قدیمی + HTML جدید: چاپ واقعی کار نمی‌کند تا همین پوسته ویندوز نصب شود. HTML جدید PrintDocument/PrintAsync می‌خواهد؛ exe قدیمی printto جعلی دارد.",
            "این محیط لینوکس است؛ PrintAsync و اسپولر اینجا وجود ندارد.",
            "PRINT_SUBMITTED به‌معنی برگه روی میز نیست. گیر کردن کاغذ، چاپگر خاموش بعد از پذیرش صف، یا درایور خراب را برنامه نمی‌تواند از اینجا ببیند.",
            "WebView2 Runtime روی سیستم فروشگاه باید نصب باشد.",
            "کپی‌های قدیمی desktop/Sirman_Install_Kit و Sirman_Windows_Install هنوز سورس printto دارند؛ منبع ساخت کیت desktop/Sirman.Desktop است.",
        ],
    )

    add_h(doc, "۱۲) آنچه در کد اصلاح شد")
    add_bullets(
        doc,
        [
            "حذف Process.Start / printto به‌عنوان مسیر چاپ.",
            "چاپ از WebView2.PrintAsync به اسپولر ویندوز.",
            "حذف ساخت چاپگر جعلی Microsoft Print to PDF.",
            "pcDoPrint دیگر نمونه پیش‌نمایش را چاپ نمی‌کند؛ بدون سند زنده NO_DOCUMENT می‌دهد.",
            "چاپ سریع فاکتور و فروش از همان موتور می‌گذرد و InvoiceId/saleUid را می‌فرستد.",
            "پیام خطا به‌جای «ارسال شد» وقتی چاپ شکست می‌خورد.",
        ],
    )

    add_h(doc, "۱۳) حکم نهایی")
    add_p(doc, STATUS, bold=True)
    add_p(
        doc,
        "برای خارج شدن از BLOCKED باید روی رایانه ویندوز فروشگاه با کیت 1405.5.25β این کارها انجام شود: Sirman.exe همین نسخه باز شود، چاپگر واقعی در مرکز پرینت دیده شود، یک فاکتور زنده چاپ شود، و برگه کاغذ با همان شماره/شناسه بیرون بیاید. تا آن لحظه چاپ فیزیکی تأییدشده نیست.",
    )
    add_p(doc, "کیت: Sirman_Setup_1405.5.25β.zip")
    add_p(doc, "HTML: Sirman_Final.html و Sirman_Final_1405.5.25β.html")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    for c in COPIES:
        c.parent.mkdir(parents=True, exist_ok=True)
        doc.save(c)
    print(OUT, OUT.stat().st_size)


if __name__ == "__main__":
    main()
