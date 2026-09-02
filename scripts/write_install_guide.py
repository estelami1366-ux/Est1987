# -*- coding: utf-8 -*-
"""Generate the from-scratch Sirman install guide (DOCX + TXT) from SIRMAN_VERSION.json."""
from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

ROOT = Path("/workspace")
VER = json.loads((ROOT / "SIRMAN_VERSION.json").read_text(encoding="utf-8"))
VERSION = VER["appFa"]
VERSION_LATIN = VER["app"]
DATE_FA = "۱۴۰۵/۰۶/۰۳"
DATE_LATIN = VER["date"]
ZIP_NAME = f"Sirman_Setup_{VERSION_LATIN}.zip"

OUT_DOCX = ROOT / "راهنمای_نصب_و_آپدیت.docx"
OUT_TXT = ROOT / "راهنمای_نصب_از_صفر.txt"
COPIES = [
    ROOT / "desktop" / "Sirman_Install_Kit",
    ROOT / "desktop" / "Sirman_Windows_Install",
    ROOT / "deliveries" / "Reports",
    Path("/opt/cursor/artifacts"),
]

WARN = RGBColor(0xB4, 0x53, 0x09)

TXT = f"""راهنمای نصب سیرمان
لایق الکترونیک پارسیان
نسخه: {VERSION}   ({VERSION_LATIN})
تاریخ: {DATE_FA}

این فایل راهنمای نصب فروشگاه است. لازم نیست برنامه‌نویس باشید.


════════════════════════════════════
۱) برنامه کدام فایل است؟
════════════════════════════════════

فایل نصب فروشگاه این است:

  {ZIP_NAME}

این یک فایل فشرده است. برنامه داخل آن است.

این‌ها برنامه نیستند:
  • فایل JSON حدود ۱ کیلوبایت داخل پوشه updates
  • پیام چت یا متن راهنما
  • فایل گزارش .md

فایل یک‌کیلوبایتی برنامه نیست.
اگر فقط یک JSON کوچک کپی کردید، هیچ برنامه‌ای نصب نشده است.

برنامه واقعی بعد از نصب این دو فایل کنار هم است:
  Sirman.exe
  Sirman_Final.html   (حدود ۱٫۶ مگابایت، نه چند کیلوبایت)


════════════════════════════════════
۲) پیش‌نیاز ویندوز فروشگاه
════════════════════════════════════

روی همان کامپیوتر فروشگاه لازم است:

  • ویندوز ۱۰ یا ۱۱، ۶۴ بیتی
  • این کیت خودکفا است؛ Runtime دات‌نت داخل پوشه App است
    لازم نیست جداگانه .NET نصب کنید
  • WebView2
    معمولاً با Microsoft Edge هست.
    اگر exe باز شد و صفحه سفید بود، WebView2 را نصب کنید:
    https://developer.microsoft.com/microsoft-edge/webview2/


════════════════════════════════════
۳) نصب از صفر — روش درست
════════════════════════════════════

فایل‌ها را یکی‌یکی کپی نکنید. از نصب.bat استفاده کنید.

گام ۱
  فایل {ZIP_NAME} را روی کامپیوتر فروشگاه کپی کنید.
  روی فلش قفل‌شده اجرا نکنید؛ اول روی هارد کپی کنید.
  مثال: Documents یا Desktop.

گام ۲
  روی zip راست‌کلیک → Extract All / استخراج همه.
  یک پوشه جدید ساخته می‌شود.

گام ۳
  داخل پوشه استخراج‌شده این‌ها را باید ببینید:
    نصب.bat
    SETUP.bat
    App\\Sirman.exe
    App\\Sirman_Final.html
    همین فایل ورد راهنما

گام ۴
  فقط روی «نصب.bat» دوبار کلیک کنید.
  (اگر ویندوز هشدار داد، Run anyway بزنید.)

گام ۵
  پوشه نصب را انتخاب کنید.
  پیشنهاد: Documents\\Sirman
  میانبر دسکتاپ را تأیید کنید.

گام ۶
  از آیکون دسکتاپ یا منوی Start برنامه «سیرمان» را باز کنید.

گام ۷
  پایین سایدبار سمت راست نسخه را ببینید.
  باید باشد: {VERSION}

تمام. نصب شد.


════════════════════════════════════
۴) اگر Sirman.exe باز نشد
════════════════════════════════════

• پنجره باز می‌شود ولی صفحه سفید است:
  WebView2 را نصب کنید.
  Sirman_Final.html باید کنار Sirman.exe باشد.

• می‌خواهید بدون exe کار کنید:
  داخل پوشه نصب، Sirman_Start.bat را بزنید.
  پنجره مشکی را نبندید.


════════════════════════════════════
۵) داده کجاست؟ خیلی مهم
════════════════════════════════════

فاکتور، گارانتی، انبار و مخاطب داخل فایل HTML یا zip نیست.
روی همان کامپیوتر، در حافظه همین برنامه ذخیره می‌شود.

  • اگر فقط فایل برنامه را به کامپیوتر دیگر ببرید، آنجا خالی باز می‌شود.
  • قبل از تعویض سیستم یا نصب دوباره، از داخل برنامه
    «ورود/خروج داده» بک‌آپ بگیرید.
  • روی سیستم جدید همان بک‌آپ را با حالت «جایگزینی» برگردانید.


════════════════════════════════════
۶) اگر از قبل برنامه را دارید (آپدیت، نه نصب اول)
════════════════════════════════════

۱) از «ورود/خروج داده» بک‌آپ کامل بگیرید.
۲) همین کیت zip جدید را Extract کنید و دوباره نصب.bat را بزنید
   روی همان پوشه نصب قبلی (مثلاً Documents\\Sirman).
۳) کیت، فایل Pending قدیمی ۱ کیلوبایتی را با آپدیت کامل همین نسخه عوض می‌کند.
۴) برنامه را باز کنید. نسخه باید {VERSION} باشد.
۵) لیست فاکتور و گارانتی باید همان قبلی باشد.

فایل JSON یک‌کیلوبایتی را به‌جای برنامه بارگذاری نکنید.


════════════════════════════════════
۷) چک درست نصب شدن
════════════════════════════════════

☐ zip استخراج شده و نصب.bat زده شده
☐ Sirman.exe و Sirman_Final.html (حدود ۱٫۶ مگابایت) کنار هم هستند
☐ برنامه باز می‌شود
☐ پایین سایدبار نسخه {VERSION} است
☐ اگر داده قبلی داشتید، فاکتور و گارانتی هنوز هستند


════════════════════════════════════
۸) حذف
════════════════════════════════════

منوی Start → Programs → Sirman → Uninstall SIRMAN
یا فایل Uninstall-Sirman.bat داخل پوشه نصب.

حذف سطح ۱ میانبر و برنامه را برمی‌دارد؛ بک‌آپ و WebView2 را پاک نمی‌کند.
پاک‌سازی کامل سطح ۲ جداست و فقط با تایپ «تایید» داده را حذف می‌کند.
"""


def set_run_rtl(run):
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)
    cs = OxmlElement("w:cs")
    rPr.append(cs)


def set_para_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in paragraph.runs:
        set_run_rtl(run)
        run.font.name = "Tahoma"
        run._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
        run.font.size = Pt(12)


def add_p(doc, text, style="Normal", bold=False, color=None, size=12):
    p = doc.add_paragraph()
    if style == "Heading 1":
        p.style = doc.styles["Heading 1"]
        size = 22
        bold = True
    elif style == "Heading 2":
        p.style = doc.styles["Heading 2"]
        size = 14
        bold = True
    elif style == "Heading 3":
        p.style = doc.styles["Heading 3"]
        size = 13
        bold = True
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Tahoma"
    if color:
        run.font.color.rgb = color
    set_para_rtl(p)
    return p


def add_bullets(doc, items, bold=False, color=None):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        set_para_rtl(p)
        for run in p.runs:
            run.font.name = "Tahoma"
            run.font.size = Pt(12)
            run.bold = bold
            if color:
                run.font.color.rgb = color


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.right_margin = Cm(2)
    section.left_margin = Cm(2)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    section._sectPr.append(bidi)

    add_p(doc, "راهنمای نصب سیرمان", "Heading 1")
    add_p(doc, "لایق الکترونیک پارسیان")
    add_p(doc, f"نسخه نرم‌افزار: {VERSION}   ({VERSION_LATIN})", bold=True)
    add_p(doc, f"تاریخ: {DATE_FA}")
    add_p(doc, "این راهنما برای نصب روی کامپیوتر فروشگاه است. از اول تا آخر پیش بروید.")

    add_p(doc, "۱) برنامه کدام فایل است؟", "Heading 2")
    add_p(doc, f"فایل نصب فروشگاه:  {ZIP_NAME}", bold=True)
    add_p(doc, "این یک فایل فشرده است. برنامه داخل آن است.")
    add_p(doc, "این‌ها برنامه نیستند:", bold=True, color=WARN)
    add_bullets(
        doc,
        [
        "فایل JSON حدود ۱ کیلوبایت داخل پوشه updates",
        "پیام چت یا متن راهنما",
        "فایل گزارش .md",
    ],
    color=WARN,
)
    add_p(doc, "فایل یک‌کیلوبایتی برنامه نیست.", bold=True, color=WARN)
    add_p(
        doc,
        "اگر فقط یک JSON کوچک کپی کردید، هیچ برنامه‌ای نصب نشده است.",
        bold=True,
        color=WARN,
    )
    add_p(doc, "برنامه واقعی بعد از نصب این دو فایل کنار هم است:", bold=True)
    add_bullets(
        doc,
        [
            "Sirman.exe",
            "Sirman_Final.html  — حدود ۱٫۶ مگابایت. اگر چند کیلوبایت بود، فایل اشتباه است.",
        ],
    )

    add_p(doc, "۲) پیش‌نیاز ویندوز فروشگاه", "Heading 2")
    add_bullets(
        doc,
        [
            "ویندوز ۱۰ یا ۱۱، ۶۴ بیتی",
            "این کیت خودکفا است؛ Runtime دات‌نت داخل پوشه App است و لازم نیست جداگانه نصب شود",
            "WebView2 — معمولاً با Edge هست. اگر exe صفحه سفید داد، Runtime وب‌ویو را نصب کنید.",
        ],
    )

    add_p(doc, "۳) نصب از صفر — روش درست", "Heading 2")
    add_p(doc, "فایل‌ها را یکی‌یکی کپی نکنید. از نصب.bat استفاده کنید.", bold=True)
    add_p(doc, "گام ۱", "Heading 3")
    add_p(doc, f"فایل {ZIP_NAME} را روی هارد فروشگاه کپی کنید. از روی فلش قفل‌شده اجرا نکنید.")
    add_p(doc, "گام ۲", "Heading 3")
    add_p(doc, "روی zip راست‌کلیک کنید و Extract All / استخراج همه را بزنید.")
    add_p(doc, "گام ۳", "Heading 3")
    add_p(doc, "داخل پوشه باید نصب.bat و پوشه App (با Sirman.exe و Sirman_Final.html) را ببینید.")
    add_p(doc, "گام ۴", "Heading 3")
    add_p(doc, "فقط روی «نصب.bat» دوبار کلیک کنید. اگر ویندوز هشدار داد، Run anyway بزنید.")
    add_p(doc, "گام ۵", "Heading 3")
    add_p(doc, "پوشه نصب را انتخاب کنید. پیشنهاد: Documents\\Sirman. میانبر دسکتاپ را تأیید کنید.")
    add_p(doc, "گام ۶", "Heading 3")
    add_p(doc, "از آیکون دسکتاپ یا منوی Start برنامه «سیرمان» را باز کنید.")
    add_p(doc, "گام ۷", "Heading 3")
    add_p(doc, f"پایین سایدبار سمت راست نسخه باید {VERSION} باشد. تمام.")

    add_p(doc, "۴) اگر Sirman.exe باز نشد", "Heading 2")
    add_bullets(
        doc,
        [
            "صفحه سفید: WebView2 را نصب کنید. HTML باید کنار exe باشد.",
            "بدون exe: داخل پوشه نصب Sirman_Start.bat را بزنید و پنجره مشکی را نبندید.",
        ],
    )

    add_p(doc, "۵) داده کجاست؟", "Heading 2")
    add_p(
        doc,
        "فاکتور، گارانتی، انبار و مخاطب داخل zip یا HTML نیست. روی همان کامپیوتر ذخیره می‌شود.",
        bold=True,
        color=WARN,
    )
    add_bullets(
        doc,
        [
            "اگر فقط فایل برنامه را به سیستم دیگر ببرید، آنجا خالی باز می‌شود.",
            "قبل از تعویض سیستم از داخل برنامه «ورود/خروج داده» بک‌آپ بگیرید.",
            "روی سیستم جدید همان بک‌آپ را با حالت جایگزینی برگردانید.",
        ],
    )

    add_p(doc, "۶) اگر از قبل برنامه را دارید", "Heading 2")
    add_bullets(
        doc,
        [
            "اول از «ورود/خروج داده» بک‌آپ کامل بگیرید.",
            "همین zip را Extract کنید و دوباره نصب.bat را روی همان پوشه نصب قبلی بزنید.",
            "کیت، فایل Pending قدیمی ۱ کیلوبایتی را با آپدیت کامل همین نسخه عوض می‌کند.",
            f"بعد از باز کردن، نسخه باید {VERSION} باشد و فاکتور قبلی سر جایش.",
        ],
    )
    add_p(
        doc,
        "فایل JSON یک‌کیلوبایتی را به‌جای برنامه بارگذاری نکنید.",
        bold=True,
        color=WARN,
    )

    add_p(doc, "۷) چک درست نصب شدن", "Heading 2")
    add_bullets(
        doc,
        [
            "zip استخراج شده و نصب.bat اجرا شده است.",
            "Sirman.exe و Sirman_Final.html حدود ۱٫۶ مگابایت کنار هم هستند.",
            "برنامه باز می‌شود.",
            f"پایین سایدبار نسخه {VERSION} است.",
            "اگر داده قبلی داشتید، فاکتور و گارانتی هنوز هستند.",
        ],
    )

    add_p(doc, "۸) حذف", "Heading 2")
    add_p(doc, "منوی Start → Programs → Sirman → Uninstall SIRMAN، یا Uninstall-Sirman.bat داخل پوشه نصب.")
    add_p(doc, "حذف سطح ۱ برنامه و میانبر را برمی‌دارد؛ WebView2 و بک‌آپ را پاک نمی‌کند. پاک‌سازی کامل سطح ۲ جداست و فقط با تایپ «تایید» داده را حذف می‌کند.")

    doc.save(str(OUT_DOCX))


def main() -> None:
    OUT_TXT.write_text(TXT, encoding="utf-8-sig")
    build_docx()
    payload_docx = OUT_DOCX.read_bytes()
    payload_txt = OUT_TXT.read_bytes()
    for folder in COPIES:
        folder.mkdir(parents=True, exist_ok=True)
        (folder / OUT_DOCX.name).write_bytes(payload_docx)
        (folder / OUT_TXT.name).write_bytes(payload_txt)
    print("wrote", OUT_DOCX, OUT_DOCX.stat().st_size)
    print("wrote", OUT_TXT, OUT_TXT.stat().st_size)


if __name__ == "__main__":
    main()
