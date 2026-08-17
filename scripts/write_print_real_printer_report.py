# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

OUT = Path("/workspace/PRINT_CENTER_REAL_PRINTER_FIX_REPORT.docx")
COPIES = [Path("/workspace/deliveries/Reports/PRINT_CENTER_REAL_PRINTER_FIX_REPORT.docx")]
STATUS = "BLOCKED — PRINTER ENVIRONMENT ISSUE"


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        from docx.oxml import OxmlElement
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
    bidi.set(qn("w:val"), "1")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_h(doc, text, size=16, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "Tahoma"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    if color:
        run.font.color.rgb = RGBColor(*color)
    set_rtl(p)


def add_p(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    run.font.name = "Tahoma"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    set_rtl(p)


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        run.font.size = Pt(11)
        run.font.name = "Tahoma"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
        set_rtl(p)


def main():
    doc = Document()
    for s in doc.sections:
        s.right_margin = Cm(2)
        s.left_margin = Cm(2)

    add_h(doc, "گزارش اصلاح اتصال چاپگر واقعی ویندوز", 20)
    add_p(doc, "نسخه: ۱۴۰۵.۵.۲۶α / 1405.5.26.1 — تاریخ ۱۴۰۵/۰۵/۲۶")
    add_h(doc, STATUS, 16, color=(176, 32, 32))
    add_p(doc, "FIXED گفته نمی‌شود: این محیط لینوکس است و برگه کاغذ چاپ نشد. خروجی PDF موفقیت چاپ نیست.", bold=True)

    add_h(doc, "۱) نقطه شکست دقیق")
    add_p(doc, "شکست روی کاغذ نبود؛ کار چاپ به چاپگر فایل ویندوز می‌رفت.")
    add_bullets(doc, [
        "GetPrinters چاپگرهای نصب‌شده از جمله Microsoft Print to PDF را برمی‌گرداند (این چاپگر جعلی ساخته نمی‌شد؛ روی ویندوز واقعاً نصب است).",
        "مرکز پرینت چاپگر isDefault را انتخاب می‌کرد. روی بیشتر ویندوزها پیش‌فرض Microsoft Print to PDF است.",
        "ResolvePrinter اگر نام خالی یا «PDF» بود، همان پیش‌فرض ویندوز را برمی‌گرداند — یعنی PDF.",
        "دکمه ذخیره PDF نام چاپگر را روی Microsoft Print to PDF می‌نوشت؛ چاپ بعدی همان را استفاده می‌کرد.",
        "PrintAsync با PrinterName=Microsoft Print to PDF فایل PDF می‌ساخت و PRINT_SUBMITTED می‌داد.",
        "نتیجه تست فروشگاه: سند ساخته می‌شد، PDF بیرون می‌آمد، صف چاپگر واقعی خالی می‌ماند.",
    ])

    add_h(doc, "۲) علت ریشه‌ای")
    add_p(doc, "مسیر PRINT و مسیر PDF یکی بود. پیش‌فرض ویندوز / پروفایل PDF / ذخیره PDF، PrinterName را به چاپگر فایل می‌رساند. WebView2.PrintAsync همان نام را به اسپولر می‌داد و اسپولر PDF می‌ساخت. این چاپ کاغذ نیست.")

    add_h(doc, "۳) جریان چاپ فعلی")
    add_p(doc, "چاپ: دکمه چاپ → printEnginePrintHtml(purpose=print) → رد چاپگر PDF → sirmanHost.PrintDocument → WindowsPrintHost → Resolve فقط چاپگر فیزیکی → OpenPrinter → PrintSettings.PrinterName=نام دقیق → PrintAsync → PRINT_SUBMITTED فقط اگر چاپگر فیزیکی باشد.")
    add_p(doc, "PDF: دکمه ذخیره PDF جدا → purpose=pdf → به مسیر چاپ کاغذ نمی‌رود و چاپگر چاپ را عوض نمی‌کند.")

    add_h(doc, "۴) کشف چاپگر")
    add_p(doc, "منبع: PrinterSettings.InstalledPrinters. فیلدها: name, isDefault, isValid, isPhysical, isPdf, kind. چاپگر جعلی ساخته نمی‌شود. defaultPhysicalPrinter جدا از defaultPrinter است.")

    add_h(doc, "۵) انتخاب چاپگر")
    add_p(doc, "UI مقدار option را همان name ویندوز می‌گذارد. اگر UI «Microsoft Print to PDF» بفرستد، JS قبل از Host با PDF_NOT_PRINT قطع می‌کند. Host هم اگر برسد رد می‌کند. Match سه‌طرفه برای چاپگر فیزیکی لازم است.")

    add_h(doc, "۶) PrintSettings")
    add_p(doc, "بعد از Resolve، settings.PrinterName همان نام InstalledPrinters است. کاغذ A4/A5/80mm/برچسب، جهت، کپی و حاشیه ست می‌شود. اگر PrinterName خالی بماند PRINT_ASYNC_FAILED.")

    add_h(doc, "۷) WebView2 PrintAsync")
    add_p(doc, "PrintAsync با همان settings اجرا می‌شود. Succeeded روی چاپگر فیزیکی → PRINT_SUBMITTED. PrinterUnavailable → PRINTER_UNAVAILABLE. OtherError/Exception → PRINT_ASYNC_FAILED با متن خطا. WebView آماده نشد → PRINT_WEBVIEW_FAILED. try/catch خالی نیست.")

    add_h(doc, "۸) اسپولر ویندوز")
    add_p(doc, "قبل از PrintAsync، OpenPrinter روی همان نام زده می‌شود. شکست OpenPrinter → PRINT_SPOOLER_FAILED. لاگ JSONL: PRINT_REQUESTED, PRINTER_RESOLVED, PRINT_SETTINGS_CREATED, PRINT_ASYNC_STARTED, PRINT_ASYNC_COMPLETED, PRINT_SUBMITTED / PRINT_FAILED. مسیر: %AppData%\\Sirman\\print\\print-jobs.jsonl")

    add_h(doc, "۹) دلیل PDF")
    add_p(doc, "PDF از Print to PDF ویندوز می‌آمد، نه از یک Export جدا. این رفتار از مسیر PRINT حذف شد. دکمه ذخیره PDF جدا ماند.")

    add_h(doc, "۱۰) اصلاح")
    add_bullets(doc, [
        "چاپگر مجازی (PDF/XPS/Fax/OneNote) برای PRINT رد می‌شود.",
        "پیش‌فرض خالی دیگر به PDF نمی‌افتد؛ اگر فقط PDF باشد PDF_NOT_PRINT.",
        "ذخیره PDF چاپگر چاپ را عوض نمی‌کند.",
        "UI چاپگر فایل را «خروجی فایل — چاپ نیست» نشان می‌دهد و پیش‌فرض انتخابی فقط فیزیکی است.",
        "لاگ مرحله‌ای و کدهای PRINT_ASYNC_FAILED / PRINT_WEBVIEW_FAILED / PRINT_SPOOLER_FAILED / PRINT_TIMEOUT / PDF_NOT_PRINT.",
    ])

    add_h(doc, "۱۱) تست ویندوز واقعی")
    add_p(doc, "این عامل ابری لینوکس است. Printers & scanners، صفحه تست ویندوز، Sirman.exe، صف چاپ و برگه کاغذ اینجا اجرا نشد.")
    add_p(doc, "تست Node: اگر Printer=Microsoft Print to PDF باشد Host صدا زده نمی‌شود و PDF_NOT_PRINT است. اگر HP باشد purpose=print و نام HP فرستاده می‌شود. تست HTML ۵۳۸+ و C# ۹۲ در انتهای کار ثبت می‌شود.")

    add_h(doc, "۱۲) نتیجه چاپ فیزیکی")
    add_p(doc, "برگه کاغذ در این محیط چاپ نشد. اگر صفحه تست ویندوز چاپ شود ولی سیرمان نه، مشکل در Sirman/WebView2/PrintSettings است نه درایور؛ آن را فقط روی PC فروشگاه با همین EXE نسخه ۱۴۰۵.۵.۲۶α می‌توان دید.")

    add_h(doc, "۱۳) باقی‌مانده")
    add_bullets(doc, [
        "باید EXE همین نسخه نصب شود؛ HTML جدید روی EXE قدیمی کافی نیست.",
        "چاپگر پیش‌فرض ویندوز اگر PDF باشد دیگر برای چاپ انتخاب نمی‌شود؛ کاربر باید چاپگر HP/Canon/... را در مرکز پرینت ببیند و انتخاب کند.",
        "PRINT_SUBMITTED یعنی اسپولر پذیرفت، نه اینکه کاغذ روی میز دیده شده.",
        "تأیید RTL/A4 روی کاغذ فقط روی ویندوز فروشگاه.",
    ])

    add_h(doc, "حکم")
    add_p(doc, STATUS, bold=True)
    add_p(doc, "FIXED — REAL PRINTER PRINT VERIFIED فقط وقتی یک فاکتور زنده با همین Sirman.exe روی چاپگر واقعی کاغذ بدهد.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    for c in COPIES:
        c.parent.mkdir(parents=True, exist_ok=True)
        doc.save(c)
    print(OUT, OUT.stat().st_size)


if __name__ == "__main__":
    main()
